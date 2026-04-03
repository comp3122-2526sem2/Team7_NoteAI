import io
import logging
import os
import uuid
from pathlib import Path

from fastapi import APIRouter, BackgroundTasks, HTTPException, UploadFile, status
from sqlalchemy import select

from database import SessionLocal
from deps import AdminUser, CurrentUser, DbDep, TeacherUser
from models import (
    Course, CourseStudent, CourseTeacher, Document, StudentUser,
    SyllabusGenerationPrompt,
    TeacherUser as TeacherModel, UserRole,
    DEFAULT_SYLLABUS_GENERATION_PROMPT,
)
from models.document import ConversionStatus, DocumentType
from openai_client import chat_complete
from schemas import (
    AssignTeacherRequest,
    CourseCreate,
    CourseOut,
    CourseUpdate,
    EnrollStudentRequest,
    SyllabusUploadOut,
    UserOut,
)

logger = logging.getLogger(__name__)

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", str(Path(__file__).parent.parent / "uploads")))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

SYLLABUS_ALLOWED_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "text/plain": "txt",
    "text/markdown": "md",
}


def _extract_text(file_bytes: bytes, content_type: str) -> str:
    """Extract plain text from an uploaded file based on its MIME type."""
    if content_type in ("text/plain", "text/markdown"):
        return file_bytes.decode("utf-8", errors="replace")

    if content_type == "application/pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    if content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        return "\n".join(para.text for para in doc.paragraphs)

    raise ValueError(f"Cannot extract text from '{content_type}'")

router = APIRouter(prefix="/courses", tags=["Courses"])


def _get_course_or_404(course_id: uuid.UUID, db) -> Course:
    course = db.get(Course, course_id)
    if not course:
        raise HTTPException(status_code=404, detail="Course not found.")
    return course


@router.get("", response_model=list[CourseOut])
def list_courses(current_user: CurrentUser, db: DbDep):
    if current_user.role == UserRole.admin:
        return db.scalars(select(Course).order_by(Course.created_at.desc())).all()
    if current_user.role == UserRole.teacher:
        return db.scalars(
            select(Course)
            .join(CourseTeacher, CourseTeacher.course_id == Course.id)
            .where(CourseTeacher.teacher_id == current_user.id)
        ).all()
    # student
    return db.scalars(
        select(Course)
        .join(CourseStudent, CourseStudent.course_id == Course.id)
        .where(CourseStudent.student_id == current_user.id)
    ).all()


@router.post("", response_model=CourseOut, status_code=status.HTTP_201_CREATED)
def create_course(body: CourseCreate, current_user: TeacherUser, db: DbDep):
    course = Course(**body.model_dump())
    db.add(course)
    db.flush()
    db.add(CourseTeacher(teacher_id=current_user.id, course_id=course.id))
    db.commit()
    db.refresh(course)
    return course


@router.get("/{course_id}", response_model=CourseOut)
def get_course(course_id: uuid.UUID, _: CurrentUser, db: DbDep):
    return _get_course_or_404(course_id, db)


@router.put("/{course_id}", response_model=CourseOut)
def update_course(course_id: uuid.UUID, body: CourseUpdate, _: TeacherUser, db: DbDep):
    course = _get_course_or_404(course_id, db)
    for field, value in body.model_dump(exclude_none=True).items():
        setattr(course, field, value)
    db.commit()
    db.refresh(course)
    return course


@router.delete("/{course_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_course(course_id: uuid.UUID, _: AdminUser, db: DbDep):
    course = _get_course_or_404(course_id, db)
    db.delete(course)
    db.commit()


# ── Students ──────────────────────────────────────────────────────────────────

@router.get("/{course_id}/students", response_model=list[UserOut])
def list_students(course_id: uuid.UUID, _: TeacherUser, db: DbDep):
    _get_course_or_404(course_id, db)
    enrollments = db.scalars(
        select(CourseStudent).where(CourseStudent.course_id == course_id)
    ).all()
    student_ids = [e.student_id for e in enrollments]
    students = [db.get(StudentUser, sid) for sid in student_ids]
    return [s.user for s in students if s]


@router.post("/{course_id}/students", status_code=status.HTTP_201_CREATED)
def enroll_student(course_id: uuid.UUID, body: EnrollStudentRequest, _: TeacherUser, db: DbDep):
    _get_course_or_404(course_id, db)
    student = db.scalar(select(StudentUser).where(StudentUser.student_id == body.student_id))
    if not student:
        raise HTTPException(status_code=404, detail="Student not found.")
    existing = db.scalar(
        select(CourseStudent).where(
            CourseStudent.course_id == course_id,
            CourseStudent.student_id == student.id,
        )
    )
    if existing:
        raise HTTPException(status_code=409, detail="Student already enrolled.")
    db.add(CourseStudent(course_id=course_id, student_id=student.id))
    db.commit()
    return {"detail": "Student enrolled."}


@router.delete("/{course_id}/students/{student_id}", status_code=status.HTTP_204_NO_CONTENT)
def unenroll_student(course_id: uuid.UUID, student_id: uuid.UUID, _: TeacherUser, db: DbDep):
    enrollment = db.scalar(
        select(CourseStudent).where(
            CourseStudent.course_id == course_id,
            CourseStudent.student_id == student_id,
        )
    )
    if not enrollment:
        raise HTTPException(status_code=404, detail="Enrollment not found.")
    db.delete(enrollment)
    db.commit()


# ── Teachers ──────────────────────────────────────────────────────────────────

@router.get("/{course_id}/teachers", response_model=list[UserOut])
def list_teachers(course_id: uuid.UUID, _: CurrentUser, db: DbDep):
    _get_course_or_404(course_id, db)
    assignments = db.scalars(
        select(CourseTeacher).where(CourseTeacher.course_id == course_id)
    ).all()
    teacher_ids = [a.teacher_id for a in assignments]
    teachers = [db.get(TeacherModel, tid) for tid in teacher_ids]
    return [t.user for t in teachers if t]


@router.post("/{course_id}/teachers", status_code=status.HTTP_201_CREATED)
def assign_teacher(course_id: uuid.UUID, body: AssignTeacherRequest, _: AdminUser, db: DbDep):
    _get_course_or_404(course_id, db)
    teacher = db.scalar(select(TeacherModel).where(TeacherModel.teacher_id == body.teacher_id))
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found.")
    existing = db.scalar(
        select(CourseTeacher).where(
            CourseTeacher.course_id == course_id,
            CourseTeacher.teacher_id == teacher.id,
        )
    )
    if existing:
        raise HTTPException(status_code=409, detail="Teacher already assigned.")
    db.add(CourseTeacher(course_id=course_id, teacher_id=teacher.id))
    db.commit()
    return {"detail": "Teacher assigned."}


@router.delete("/{course_id}/teachers/{teacher_id}", status_code=status.HTTP_204_NO_CONTENT)
def remove_teacher(course_id: uuid.UUID, teacher_id: uuid.UUID, _: AdminUser, db: DbDep):
    assignment = db.scalar(
        select(CourseTeacher).where(
            CourseTeacher.course_id == course_id,
            CourseTeacher.teacher_id == teacher_id,
        )
    )
    if not assignment:
        raise HTTPException(status_code=404, detail="Teacher assignment not found.")
    db.delete(assignment)
    db.commit()


# ── Syllabus helpers ───────────────────────────────────────────────────────────

# GitHub Models free tier caps input at ~8 000 tokens per request.
# Budget breakdown (chars ≈ tokens × 4):
#   prompt template overhead  ~  1 600 chars (400 tokens)
#   system message            ~    200 chars ( 50 tokens)
#   expected syllabus output  ~ 10 000 chars (2 500 tokens)
#   → safe budget for {file_content}: 10 000 chars (2 500 tokens)
#
# Chunk size is kept small so each individual summarisation call also stays
# safely within the same 8 000-token window.

_MAX_CONTENT_CHARS = 10_000   # hard ceiling for {file_content} in the final prompt
_CHUNK_SIZE = 3_500           # chars per chunk (~875 tokens)
_SUMMARY_MAX_TOKENS = 120     # cap each chunk summary to ~120 tokens (~480 chars)


async def _summarise_chunks(text: str) -> str:
    """Summarise text in fixed-size chunks and return the joined summaries."""
    chunks = [text[i: i + _CHUNK_SIZE] for i in range(0, len(text), _CHUNK_SIZE)]
    total = len(chunks)
    summaries: list[str] = []

    for idx, chunk in enumerate(chunks, start=1):
        summary = await chat_complete(
            (
                f"Briefly summarise part {idx}/{total} of a course document. "
                f"Keep only key topics, objectives, schedule items, and "
                f"assessment details. Be concise.\n\n{chunk}"
            ),
            system="You are a concise academic content summariser.",
            temperature=0.3,
            max_tokens=_SUMMARY_MAX_TOKENS,
        )
        summaries.append(summary.strip())

    return "\n\n".join(summaries)


async def _condense_for_prompt(text: str) -> str:
    """
    Map-reduce condensation for large documents.

    Pass 1 – if the raw text exceeds the safe budget, split into chunks and
             summarise each one individually (bounded by _SUMMARY_MAX_TOKENS).
    Pass 2 – if the combined summaries are still too long (unlikely but
             possible with very large documents), run a second reduction pass.
    Final  – hard-truncate to _MAX_CONTENT_CHARS as a safety net.
    """
    if len(text) <= _MAX_CONTENT_CHARS:
        return text

    # Pass 1: summarise every chunk
    condensed = await _summarise_chunks(text)

    # Pass 2: if still too large, summarise the summaries
    if len(condensed) > _MAX_CONTENT_CHARS:
        condensed = await _summarise_chunks(condensed)

    # Final safety net: hard truncate
    if len(condensed) > _MAX_CONTENT_CHARS:
        condensed = condensed[:_MAX_CONTENT_CHARS] + "\n\n[Content truncated for length]"

    return condensed


# ── Syllabus background task ───────────────────────────────────────────────────

async def _generate_syllabus_bg(
    course_id: uuid.UUID,
    doc_id: uuid.UUID,
    extracted_text: str,
    course_name: str,
    course_description: str,
    prompt_template: str,
) -> None:
    """
    Background task: condense extracted text → call AI → persist results.

    Opens its own DB session because the request session is closed by the
    time this runs.  On any failure the Document is marked as failed.
    """
    db = SessionLocal()
    try:
        file_content = await _condense_for_prompt(extracted_text)

        prompt = prompt_template.format(
            course_name=course_name,
            course_description=course_description,
            file_content=file_content,
        )
        syllabus_md = await chat_complete(
            prompt,
            system="You are an expert curriculum designer. Always respond in markdown.",
            temperature=0.4,
        )

        doc = db.get(Document, doc_id)
        course = db.get(Course, course_id)
        if doc:
            doc.conversion_status = ConversionStatus.completed
        if course:
            course.syllabus = syllabus_md
        db.commit()
        logger.info("Syllabus generated for course %s (doc %s)", course_id, doc_id)

    except Exception:
        logger.exception("Syllabus generation failed for doc %s", doc_id)
        db.rollback()
        try:
            doc = db.get(Document, doc_id)
            if doc:
                doc.conversion_status = ConversionStatus.failed
                db.commit()
        except Exception:
            db.rollback()
    finally:
        db.close()


# ── Syllabus endpoints ─────────────────────────────────────────────────────────

@router.get("/{course_id}/syllabus", tags=["Syllabus"])
def get_syllabus(course_id: uuid.UUID, _: CurrentUser, db: DbDep):
    """Return the current AI-generated syllabus for a course (plain markdown string)."""
    course = _get_course_or_404(course_id, db)
    if not course.syllabus:
        raise HTTPException(status_code=404, detail="No syllabus has been generated for this course yet.")
    return {"course_id": course_id, "syllabus": course.syllabus}


@router.post(
    "/{course_id}/syllabus/upload",
    response_model=SyllabusUploadOut,
    status_code=status.HTTP_202_ACCEPTED,
    tags=["Syllabus"],
)
async def upload_syllabus_file(
    course_id: uuid.UUID,
    file: UploadFile,
    background_tasks: BackgroundTasks,
    current_user: TeacherUser,
    db: DbDep,
):
    """
    Accept a file upload and immediately return 202 Accepted.

    Text extraction happens synchronously (fast); AI condensation and syllabus
    generation are dispatched as a background task so the client is not blocked.
    Poll GET /documents/{document_id} and check conversion_status:
      - "pending"   → still generating
      - "completed" → Course.syllabus has been updated
      - "failed"    → generation failed; the teacher may retry
    """
    course = _get_course_or_404(course_id, db)

    if file.content_type not in SYLLABUS_ALLOWED_TYPES:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Unsupported file type '{file.content_type}'. "
                f"Allowed: {', '.join(SYLLABUS_ALLOWED_TYPES.keys())}"
            ),
        )

    file_bytes = await file.read()
    ext = SYLLABUS_ALLOWED_TYPES[file.content_type]
    saved_name = f"{uuid.uuid4()}.{ext}"
    save_path = UPLOAD_DIR / saved_name
    save_path.write_bytes(file_bytes)

    try:
        extracted_text = _extract_text(file_bytes, file.content_type)
    except Exception as exc:
        save_path.unlink(missing_ok=True)
        raise HTTPException(status_code=422, detail=f"Could not extract text from file: {exc}")

    if not extracted_text.strip():
        save_path.unlink(missing_ok=True)
        raise HTTPException(status_code=422, detail="No readable text found in the uploaded file.")

    prompt_row = db.scalar(select(SyllabusGenerationPrompt))
    prompt_template = prompt_row.prompt if prompt_row else DEFAULT_SYLLABUS_GENERATION_PROMPT

    doc = Document(
        uploaded_by=current_user.id,
        course_id=course_id,
        document_type=DocumentType.other,
        original_filename=file.filename,
        original_file_type=ext,
        original_file_path=str(save_path),
        converted_markdown=extracted_text,
        conversion_status=ConversionStatus.pending,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    background_tasks.add_task(
        _generate_syllabus_bg,
        course_id=course_id,
        doc_id=doc.id,
        extracted_text=extracted_text,
        course_name=course.name,
        course_description=course.description or "N/A",
        prompt_template=prompt_template,
    )

    return SyllabusUploadOut(course_id=course_id, document_id=doc.id, status="pending")
